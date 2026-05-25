"""Design system for the DOCX exporter.

One canonical visual language used by both the data-driven `build_docx`
and the markdown-driven `build_docx_from_markdown` (via a generated
pandoc reference doc). Mirrors the design DNA of polished one-page
research briefs (Claude artifacts, Stripe quarterly reports, a16z
research): generous whitespace, one accent color, ink-on-white tables
with hairline borders, color-coded severity chips, KPI strip on the
cover.

Public surface:
  - palette + typography constants
  - `style_table(table, *, header_fill, zebra)` — applies the brand
    table style (white-on-ink header, hairline borders, optional zebra)
  - `add_cover_page(doc, title, subtitle, kpi_strip, tagline)`
  - `add_kpi_strip(doc, kpis)` — 3-4 big-number cards in one row
  - `add_section_header(doc, label, number)` — accent number + headline
  - `add_callout(doc, body, kind)` — boxed callout (info/warn/quote)
  - `add_quote_block(doc, quote, attribution)` — left-bar accented quote
  - `add_severity_chip_run(paragraph, severity)` — inline severity pill
  - `add_divider(doc)` — horizontal hairline
  - `make_brand_reference_docx(out_path)` — generate the reference doc
    pandoc consumes for `--reference-doc` to brand markdown→docx output
  - `DESIGN_SYSTEM_PROMPT` — the strict design instructions an LLM
    should follow when planning a brief layout

All helpers expect python-docx is installed; soft-import guards live in
the parent module.
"""
from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor, Cm, Emu


# ─── Palette ──────────────────────────────────────────────────────────────
# Single accent system. Designed to print well on white and hold up in
# Google Docs / Pages / PowerPoint after upload.
#
# Aligned with the PDF brand (docs/demo_pdf/pdf_build/header.tex, Apr
# 2026): accentblue #1F4E79, accentlight #2E75B6, softgray #595959,
# rulegray #BFBFBF, codebg #F6F8FA. Same colors across PDF/DOCX/PPTX.

ACCENT = RGBColor(0x1F, 0x4E, 0x79)     # #1F4E79 — primary brand accent (matches PDF accentblue)
ACCENT_SOFT = RGBColor(0x2E, 0x75, 0xB6) # #2E75B6 — secondary accent / link (matches PDF accentlight)
ACCENT_TINT = RGBColor(0xE6, 0xEF, 0xF7) # #E6EFF7 — pale fill behind accent text

INK = RGBColor(0x0F, 0x17, 0x2A)        # slate-900, primary text
BODY = RGBColor(0x1E, 0x29, 0x3B)       # slate-800, body text
MUTE = RGBColor(0x59, 0x59, 0x59)       # PDF softgray — captions / metadata
MUTE_LIGHT = RGBColor(0x80, 0x80, 0x80) # PDF lightgray — header/footer chrome
HAIRLINE = RGBColor(0xBF, 0xBF, 0xBF)   # PDF rulegray — table borders
CARD_FILL = RGBColor(0xF6, 0xF8, 0xFA)  # PDF codebg — callout backgrounds

PAIN_HI = RGBColor(0xDC, 0x26, 0x26)    # red-600, high-severity pain
PAIN_MED = RGBColor(0xD9, 0x77, 0x06)   # amber-600
PAIN_LO = RGBColor(0x65, 0x7B, 0x83)    # neutral
WIN = RGBColor(0x05, 0x96, 0x69)        # emerald-600

WHITE = RGBColor(0xFF, 0xFF, 0xFF)


# ─── Typography ───────────────────────────────────────────────────────────
# Poppins is the brand sans-serif (matches the PDF pipeline's `sansfont`
# variable). DejaVu Sans is the body fallback. Inter is kept as a
# secondary because most macOS / modern Windows boxes ship it.
#
# python-docx writes the font name into the docx XML; if the renderer
# can't find the font it substitutes silently. Order matters: we set the
# preferred face first so Word/Pages/Google-Docs use it when available.

FONT_SANS = "Poppins"
FONT_SANS_FALLBACK = "Inter"
FONT_SANS_FALLBACK2 = "Helvetica Neue"
FONT_BODY = "DejaVu Sans"          # PDF mainfont — wide-coverage UTF-8
FONT_BODY_FALLBACK = "Helvetica"
FONT_MONO = "DejaVu Sans Mono"     # PDF monofont — same family as body
FONT_MONO_FALLBACK = "JetBrains Mono"
FONT_MONO_FALLBACK2 = "Menlo"

SIZE_TITLE = Pt(32)
SIZE_SUBTITLE = Pt(14)
SIZE_KPI_NUM = Pt(28)
SIZE_KPI_LABEL = Pt(8.5)
SIZE_H1 = Pt(20)
SIZE_H2 = Pt(15)
SIZE_H3 = Pt(12)
SIZE_H4 = Pt(10.5)
SIZE_BODY = Pt(10.5)
SIZE_QUOTE = Pt(11)
SIZE_CAPTION = Pt(8.5)
SIZE_TABLE = Pt(9.5)
SIZE_TABLE_HDR = Pt(9)
SIZE_CODE = Pt(9)


# ─── Low-level OOXML helpers (set things python-docx doesn't expose) ─────


def _set_cell_shading(cell, color: RGBColor) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{color}")
    tc_pr.append(shd)


def _set_cell_borders(cell, *, top=None, bottom=None, left=None, right=None,
                      sz: int = 4, color: RGBColor = HAIRLINE) -> None:
    """Each side: True = draw hairline; None = leave alone; False = remove.

    `sz` is in 1/8 of a point (Word convention). 4 = 0.5pt hairline.
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for side, val in (("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)):
        if val is None:
            continue
        existing = borders.find(qn(f"w:{side}"))
        if existing is not None:
            borders.remove(existing)
        b = OxmlElement(f"w:{side}")
        if val is False:
            b.set(qn("w:val"), "nil")
        else:
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), str(sz))
            b.set(qn("w:color"), f"{color}")
        borders.append(b)


def _set_cell_margins(cell, *, top: int = 80, bottom: int = 80,
                      left: int = 100, right: int = 100) -> None:
    """Margins in twentieths of a point. 100 = 5pt."""
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.find(qn("w:tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for side, val in (("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)):
        existing = mar.find(qn(f"w:{side}"))
        if existing is not None:
            mar.remove(existing)
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(val))
        m.set(qn("w:type"), "dxa")
        mar.append(m)


def _set_paragraph_border(p, *, side: str = "left", sz: int = 16,
                          color: RGBColor = ACCENT) -> None:
    """Add a single-sided border to a paragraph (used for accent quote bars)."""
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    b = OxmlElement(f"w:{side}")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), str(sz))
    b.set(qn("w:space"), "8")
    b.set(qn("w:color"), f"{color}")
    pBdr.append(b)


def _set_paragraph_shading(p, color: RGBColor) -> None:
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{color}")
    pPr.append(shd)


# ─── Run helpers ──────────────────────────────────────────────────────────


def _run(p, text: str, *, font: str = FONT_SANS, size=SIZE_BODY,
         color: RGBColor = BODY, bold: bool = False, italic: bool = False,
         small_caps: bool = False, all_caps: bool = False, spacing: int | None = None):
    r = p.add_run(text)
    r.font.name = font
    # docx requires explicit east-asian font setting too for full font swap.
    r.font.size = size
    r.font.color.rgb = color
    r.bold = bold
    r.italic = italic
    if small_caps:
        rPr = r._r.get_or_add_rPr()
        sc = OxmlElement("w:smallCaps")
        sc.set(qn("w:val"), "true")
        rPr.append(sc)
    if all_caps:
        rPr = r._r.get_or_add_rPr()
        cap = OxmlElement("w:caps")
        cap.set(qn("w:val"), "true")
        rPr.append(cap)
    if spacing is not None:
        rPr = r._r.get_or_add_rPr()
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(spacing))  # 1/20 of a point
        rPr.append(sp)
    return r


def _para(doc, *, before: int = 0, after: int = 0, alignment=None,
          line_spacing: float | None = None):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    if alignment is not None:
        p.alignment = alignment
    if line_spacing:
        fmt.line_spacing = line_spacing
        fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    return p


# ─── Table style ──────────────────────────────────────────────────────────


def style_table(table, *, header_fill: RGBColor = INK, header_text: RGBColor = WHITE,
                zebra: bool = True, numeric_cols: list[int] | None = None) -> None:
    """Apply the brand table style.

    - Header row: bold white text on Ink fill, 9pt
    - Body rows: 9.5pt, hairline borders only on top/bottom of each row
    - Optional zebra (slate-50) on alternating rows for readability
    - Numeric columns right-aligned (pass column indices)
    """
    numeric_cols = numeric_cols or []
    # Disable Word's default table style which adds heavy borders.
    table.style = None
    rows = list(table.rows)
    for ri, row in enumerate(rows):
        is_header = ri == 0
        for ci, cell in enumerate(row.cells):
            _set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # Borders: hairlines, header gets a heavier bottom rule.
            _set_cell_borders(
                cell,
                top=True if ri == 0 else False,
                bottom=True,
                left=False,
                right=False,
                sz=8 if is_header else 4,
                color=INK if is_header else HAIRLINE,
            )
            if is_header:
                _set_cell_shading(cell, header_fill)
            elif zebra and ri % 2 == 0:
                _set_cell_shading(cell, CARD_FILL)
            # Re-style every paragraph in the cell uniformly.
            for p in cell.paragraphs:
                fmt = p.paragraph_format
                fmt.space_before = Pt(0)
                fmt.space_after = Pt(0)
                if ci in numeric_cols:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                # If the cell already has runs (we added text via .text=),
                # restyle them; else leave it for caller to populate.
                for r in p.runs:
                    r.font.name = FONT_SANS
                    r.font.size = SIZE_TABLE_HDR if is_header else SIZE_TABLE
                    r.font.color.rgb = header_text if is_header else BODY
                    r.bold = bool(is_header) or r.bold


# ─── Cover + KPI strip ────────────────────────────────────────────────────


def widen_quote_columns(table, *, quote_width: float = 0.50,
                        weights: dict[str, float] | None = None,
                        page_width_in: float = 6.5) -> bool:
    """Re-allocate column widths when a table has a 'Quote' header.

    Mirrors the Lua filter the PDF pipeline uses
    (`docs/demo_pdf/pdf_build/widen-quote.lua`). Citation tables shaped
    `| Source | ID | Score | Quote |` get the Quote column 50% width
    and the rest distributed by header weight (Source 1.4, ID 1.0,
    Score 0.7, otherwise 1.0).

    Returns True if a Quote column was found and widths were adjusted,
    False otherwise — caller can fall back to default sizing.
    """
    weights = weights or {"source": 1.4, "id": 1.0, "score": 0.7}
    rows = list(table.rows)
    if not rows:
        return False
    header_cells = rows[0].cells
    headers = [(c.text or "").strip().lower() for c in header_cells]
    if "quote" not in headers:
        return False

    quote_idx = headers.index("quote")
    remaining = 1.0 - quote_width
    n = len(headers)

    weight_per_col: list[float | None] = [None] * n
    total = 0.0
    for i, h in enumerate(headers):
        if i == quote_idx:
            continue
        w = weights.get(h, 1.0)
        weight_per_col[i] = w
        total += w
    if total <= 0:
        return False

    page_w = Inches(page_width_in)
    for i, col in enumerate(table.columns):
        if i == quote_idx:
            ratio = quote_width
        else:
            ratio = (weight_per_col[i] or 1.0) / total * remaining
        target = int(page_w * ratio)
        for cell in col.cells:
            cell.width = target  # python-docx allows a per-cell override
    table.autofit = False
    return True


def add_page_chrome(doc, *, header_left: str | None = None,
                    header_right: str | None = None,
                    footer_left: str | None = None,
                    footer_link: str | None = None,
                    footer_right: str | None = None,
                    show_page_numbers: bool = True) -> None:
    """Apply brand header + footer to every page.

    Mirrors the PDF page chrome (fancyhdr in header.tex):
      - thin gray rule under the header
      - light-gray sans labels on left/right of header
      - footer: link (accent), page #, brand
    Skips a side cleanly if you pass None.
    """
    section = doc.sections[0]
    header = section.header
    footer = section.footer

    def _three_col_para(p):
        # Word's built-in tabbed three-column layout: left tab + center
        # tab at midpoint + right tab at the right margin.
        pf = p.paragraph_format
        pf.tab_stops.clear_all()
        usable = section.page_width - section.left_margin - section.right_margin
        pf.tab_stops.add_tab_stop(usable // 2, alignment=1)  # center
        pf.tab_stops.add_tab_stop(usable, alignment=2)        # right
        return p

    if header_left or header_right:
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.text = ""
        _three_col_para(hp)
        if header_left:
            _run(hp, header_left, color=MUTE_LIGHT, size=Pt(9))
        hp.add_run("\t\t")
        if header_right:
            _run(hp, header_right, color=MUTE_LIGHT, size=Pt(9))

    if footer_left or footer_link or footer_right or show_page_numbers:
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.text = ""
        _three_col_para(fp)
        if footer_link:
            _run(fp, footer_link, color=ACCENT_SOFT, size=Pt(8.5))
        elif footer_left:
            _run(fp, footer_left, color=MUTE_LIGHT, size=Pt(8.5))
        fp.add_run("\t")
        if show_page_numbers:
            # Insert a Word PAGE field. python-docx doesn't expose a
            # helper, so we drop the OOXML directly.
            run = fp.add_run()
            fld = OxmlElement("w:fldSimple")
            fld.set(qn("w:instr"), "PAGE")
            run._r.append(fld)
            for r in fp.runs:
                if r is run:
                    continue  # page field has its own formatting
            # Style the page-number run via a sibling text run; field
            # formatting is rendered by Word using the surrounding run's
            # font, so we set defaults on the paragraph.
        fp.add_run("\t")
        if footer_right:
            _run(fp, footer_right, color=MUTE_LIGHT, size=Pt(8.5))


def add_cover_page(doc, *, title: str, subtitle: str | None = None,
                   tagline: str | None = None,
                   kpis: list[tuple[str, str]] | None = None) -> None:
    """Render a magazine-style cover page.

    Layout (top → bottom):
      - 1.2" top margin of whitespace
      - Tagline (uppercase tracking, mute color)
      - Title (32pt bold, ink)
      - Subtitle (14pt italic, mute)
      - 0.6" gap
      - KPI strip (3-4 big-number cards) if given
      - Page break
    """
    # Push title down a bit
    spacer = _para(doc, before=24, after=0)

    if tagline:
        p = _para(doc, after=10)
        _run(p, tagline.upper(), color=ACCENT, bold=True, size=Pt(9.5),
             all_caps=True, spacing=40)

    p = _para(doc, after=4)
    _run(p, title, font=FONT_SANS, size=SIZE_TITLE, color=INK, bold=True)
    p.paragraph_format.line_spacing = 1.1

    if subtitle:
        p = _para(doc, after=18)
        _run(p, subtitle, font=FONT_SANS, size=SIZE_SUBTITLE, color=MUTE, italic=True)

    if kpis:
        add_kpi_strip(doc, kpis)

    doc.add_page_break()


def add_kpi_strip(doc, kpis: list[tuple[str, str]]) -> None:
    """A horizontal strip of big-number cards. Each tuple = (number, label)."""
    n = len(kpis)
    if not n:
        return
    table = doc.add_table(rows=1, cols=n)
    table.autofit = False
    # Even column widths summing to roughly the printable text width.
    total_in = 6.5
    col_w = Inches(total_in / n)
    for cell, (num, label) in zip(table.rows[0].cells, kpis):
        cell.width = col_w
        _set_cell_margins(cell, top=140, bottom=160, left=160, right=160)
        _set_cell_shading(cell, CARD_FILL)
        _set_cell_borders(cell, top=False, bottom=False, left=False, right=False)
        cell.text = ""
        p_num = cell.paragraphs[0]
        p_num.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_num.paragraph_format.space_after = Pt(2)
        _run(p_num, str(num), size=SIZE_KPI_NUM, color=INK, bold=True)
        p_lab = cell.add_paragraph()
        p_lab.paragraph_format.space_after = Pt(0)
        _run(p_lab, str(label).upper(), size=SIZE_KPI_LABEL, color=MUTE,
             all_caps=True, spacing=40)
    _para(doc, after=6)


# ─── Section helpers ──────────────────────────────────────────────────────


def add_section_header(doc, label: str, *, number: str | None = None,
                       eyebrow: str | None = None) -> None:
    """Newspaper-style section header: small eyebrow + big headline + accent."""
    if eyebrow:
        p = _para(doc, before=24, after=2)
        _run(p, eyebrow.upper(), color=ACCENT, bold=True, size=Pt(9),
             all_caps=True, spacing=40)
    p = _para(doc, after=8)
    if number:
        _run(p, f"{number}  ", size=SIZE_H1, color=ACCENT, bold=True)
    _run(p, label, size=SIZE_H1, color=INK, bold=True)
    _set_paragraph_border(p, side="bottom", sz=4, color=HAIRLINE)


def add_subsection_header(doc, label: str) -> None:
    p = _para(doc, before=14, after=4)
    _run(p, label, size=SIZE_H2, color=INK, bold=True)


def add_caption(doc, text: str) -> None:
    p = _para(doc, after=6)
    _run(p, text, size=SIZE_CAPTION, color=MUTE, italic=True)


def add_divider(doc) -> None:
    p = _para(doc, before=8, after=8)
    _set_paragraph_border(p, side="bottom", sz=4, color=HAIRLINE)


def add_quote_block(doc, quote: str, attribution: str | None = None) -> None:
    """Pull-quote with a colored left-edge bar."""
    p = _para(doc, before=6, after=2, line_spacing=1.35)
    _set_paragraph_border(p, side="left", sz=24, color=ACCENT)
    p.paragraph_format.left_indent = Inches(0.18)
    _run(p, "“" + quote.strip() + "”", size=SIZE_QUOTE, color=BODY, italic=True)
    if attribution:
        ap = _para(doc, after=10)
        ap.paragraph_format.left_indent = Inches(0.18)
        _run(ap, "— " + attribution, size=SIZE_CAPTION, color=MUTE)


def add_callout(doc, body: str, *, kind: str = "info", title: str | None = None) -> None:
    """Boxed callout. kind: 'info' (accent), 'warn' (red), 'win' (green)."""
    palette = {
        "info": (ACCENT, ACCENT_SOFT),
        "warn": (PAIN_HI, RGBColor(0xFE, 0xE2, 0xE2)),
        "win": (WIN, RGBColor(0xD1, 0xFA, 0xE5)),
    }
    border, fill = palette.get(kind, palette["info"])
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.5)
    _set_cell_shading(cell, fill)
    _set_cell_borders(cell, top=False, bottom=False, left=True, right=False, sz=24, color=border)
    _set_cell_margins(cell, top=140, bottom=140, left=160, right=160)
    cell.text = ""
    if title:
        p = cell.paragraphs[0]
        _run(p, title, size=Pt(10.5), color=INK, bold=True)
        p2 = cell.add_paragraph()
        _run(p2, body, size=SIZE_BODY, color=BODY)
    else:
        _run(cell.paragraphs[0], body, size=SIZE_BODY, color=BODY)
    _para(doc, after=4)


def add_severity_chip_run(paragraph, severity: str) -> None:
    """Inline pill: HIGH (red), MED (amber), LOW (mute)."""
    severity = (severity or "").lower()
    if severity in ("high", "hi", "critical"):
        color = PAIN_HI
        text = "HIGH"
    elif severity in ("med", "medium"):
        color = PAIN_MED
        text = "MED"
    else:
        color = PAIN_LO
        text = "LOW"
    r = paragraph.add_run("  " + text + "  ")
    r.font.name = FONT_SANS
    r.font.size = Pt(8)
    r.font.color.rgb = WHITE
    r.bold = True
    rPr = r._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{color}")
    rPr.append(shd)
    cap = OxmlElement("w:caps")
    cap.set(qn("w:val"), "true")
    rPr.append(cap)


def add_painpoint_card(doc, *, number: int, label: str, severity: str,
                       frequency: int, opportunity: float, evidence: str | None,
                       citations: list[dict] | None) -> None:
    """A polished painpoint card.

    [N. Title]                           [SEVERITY]
    Frequency 8 · Opportunity 14.0/20
    ──────────────────────────────────────────
    "evidence quote"
                          — quote attribution

    Cited posts:
      • [reddit/r/mortgages · 632▲] post title …
      • [appstore/Hearth · 1★] another …
    """
    # Title row
    p = _para(doc, before=14, after=2)
    _run(p, f"{number}. ", size=SIZE_H3, color=ACCENT, bold=True)
    _run(p, label, size=SIZE_H3, color=INK, bold=True)
    add_severity_chip_run(p, severity)

    # Metadata row
    pm = _para(doc, after=4)
    _run(pm, f"Frequency {frequency} · Opportunity {opportunity:.1f}/20",
         size=SIZE_CAPTION, color=MUTE, small_caps=True, spacing=20)
    _set_paragraph_border(pm, side="bottom", sz=4, color=HAIRLINE)

    if evidence:
        add_quote_block(doc, evidence, attribution="extracted evidence")

    if citations:
        p = _para(doc, before=4, after=2)
        _run(p, "Cited posts:", size=SIZE_CAPTION, color=INK, bold=True,
             small_caps=True, spacing=20)
        for c in citations:
            bp = _para(doc, after=0)
            bp.paragraph_format.left_indent = Inches(0.2)
            _run(bp, "• ", color=ACCENT, bold=True, size=SIZE_BODY)
            tag = f"[{c.get('source_type','?')}/{c.get('sub','?')} · {c.get('score','?')}▲]"
            _run(bp, tag + " ", size=SIZE_CAPTION, color=MUTE)
            _run(bp, (c.get("title") or "")[:140], size=SIZE_BODY, color=BODY)
            if c.get("permalink"):
                _run(bp, "  " + c["permalink"], size=SIZE_CAPTION, color=ACCENT)


# ─── Brand reference doc (for pandoc --reference-doc) ────────────────────


def make_brand_reference_docx(out_path: str | Path) -> str:
    """Generate a minimal docx whose styles pandoc copies into its output.

    Pandoc reads the named styles (Title, Subtitle, Heading 1-4, Normal,
    Block Text, Source Code, Table) from this file and applies them to
    the markdown it converts. Override the defaults so markdown→docx
    matches the same brand as the data-driven renderer.
    """
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    styles = doc.styles

    def _apply(style_name: str, *, font=FONT_SANS, size=SIZE_BODY,
               color=BODY, bold=False, italic=False,
               space_before=0, space_after=4, line_spacing=1.35):
        try:
            s = styles[style_name]
        except KeyError:
            return
        f = s.font
        f.name = font
        f.size = size
        f.color.rgb = color
        f.bold = bold
        f.italic = italic
        # East-asian font swap so the font sticks across renderers.
        rPr = s.element.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            s.element.append(rPr)
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
            rFonts.set(qn(f"w:{attr}"), font)
        # Paragraph-level
        if hasattr(s, "paragraph_format"):
            pf = s.paragraph_format
            pf.space_before = Pt(space_before)
            pf.space_after = Pt(space_after)
            pf.line_spacing = line_spacing
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    _apply("Normal", size=SIZE_BODY, color=BODY)
    _apply("Title", size=SIZE_TITLE, color=INK, bold=True,
           space_before=24, space_after=8, line_spacing=1.1)
    _apply("Subtitle", size=SIZE_SUBTITLE, color=MUTE, italic=True,
           space_before=0, space_after=18)
    _apply("Heading 1", size=SIZE_H1, color=INK, bold=True,
           space_before=24, space_after=8)
    _apply("Heading 2", size=SIZE_H2, color=INK, bold=True,
           space_before=18, space_after=6)
    _apply("Heading 3", size=SIZE_H3, color=ACCENT, bold=True,
           space_before=14, space_after=4)
    _apply("Heading 4", size=SIZE_H4, color=INK, bold=True,
           space_before=10, space_after=2)
    _apply("Quote", size=SIZE_QUOTE, color=BODY, italic=True,
           space_before=6, space_after=6)
    _apply("Intense Quote", size=SIZE_QUOTE, color=BODY, italic=True,
           space_before=6, space_after=6)

    # Add a representative title + a paragraph so pandoc's first-element
    # detection has something to anchor to.
    t = doc.add_paragraph(style="Title")
    t.add_run("Brand reference")
    sub = doc.add_paragraph(style="Subtitle")
    sub.add_run("This file's styles drive markdown → docx output.")
    doc.add_paragraph(
        "Body text uses the Normal style. Headings 1–4 carry the type ramp. "
        "Quote and Intense Quote are italic with tight leading."
    )

    doc.save(out_path)
    return str(out_path)


# ─── The strict design-rules prompt ──────────────────────────────────────
# Exposed through the MCP layer so an LLM planning a layout has the same
# constraints the renderer enforces. Keep it tight — every rule below is
# also implemented in code, so the model can't drift.

DESIGN_SYSTEM_PROMPT = """You are designing a stakeholder-ready research brief.
The renderer enforces a strict design system — your layout plan must
respect these invariants:

TYPOGRAPHY
- One sans-serif type ramp (Inter / Helvetica Neue / Calibri).
- Title 32pt · Subtitle 14pt italic · H1 20pt · H2 15pt · H3 12pt.
- Body 10.5pt @ 1.4 line-spacing. Quote 11pt italic with left accent bar.
- Code 9pt monospace on light-gray fill.

COLOR
- ONE accent (#2563EB blue). Do not introduce a second accent.
- Ink #0F172A for primary text, #1E293B for body, #64748B for captions.
- Severity chips: HIGH = #DC2626, MED = #D97706, LOW = neutral. Do not
  invent other levels.
- Hairline #E2E8F0. Card/zebra fill #F8FAFC.

LAYOUT
- Cover page = tagline (uppercase tracking) + title + subtitle + KPI strip
  of 3–4 big-number cards. Page-break before any section.
- Each section starts with an EYEBROW (uppercase accent), a numbered H1,
  and a hairline rule under it.
- Painpoint cards: number + title + severity chip + freq/opp metadata
  + quote + cited posts list. Always in this order.
- Tables: dark header row (white text on Ink), hairline-only body
  borders, optional zebra rows, right-aligned numeric columns. Never use
  the default Word "Light Grid Accent N" styles.
- Quotes: always carry an attribution (source · score · post id).

CONTENT RULES
- Every painpoint claim cites at least one post id from the corpus.
- Numbers always carry a unit (★, %, ▲, $, kbps, days, posts).
- Severity is one of {high, med, low}. No "very high", no "5/5".
- Avoid corporate filler: "leverage", "synergy", "best-in-class".
- Pull quotes are user-quoted text, never paraphrase.

LAYOUT-PLAN OUTPUT
Return a JSON object shaped like:
{
  "cover": {"title": str, "subtitle": str, "tagline": str,
            "kpis": [["1,890", "corpus posts"], ["16", "sources"], ...]},
  "sections": [
    {"number": "01", "eyebrow": "...", "label": "...", "kind": "painpoint_cards",
     "items": [...]},
    ...
  ]
}
where kind ∈ {executive_summary, painpoint_cards, competitor_matrix,
quote_wall, feature_roadmap, citation_index, callout, free_markdown}.
"""


__all__ = [
    "INK", "BODY", "MUTE", "MUTE_LIGHT", "HAIRLINE", "CARD_FILL",
    "ACCENT", "ACCENT_SOFT", "ACCENT_TINT",
    "PAIN_HI", "PAIN_MED", "PAIN_LO", "WIN", "WHITE",
    "FONT_SANS", "FONT_BODY", "FONT_MONO",
    "SIZE_TITLE", "SIZE_SUBTITLE", "SIZE_H1", "SIZE_H2", "SIZE_H3", "SIZE_H4",
    "SIZE_BODY", "SIZE_QUOTE", "SIZE_CAPTION", "SIZE_TABLE", "SIZE_TABLE_HDR",
    "SIZE_KPI_NUM", "SIZE_KPI_LABEL", "SIZE_CODE",
    "style_table", "widen_quote_columns",
    "add_cover_page", "add_kpi_strip", "add_page_chrome",
    "add_section_header", "add_subsection_header", "add_caption",
    "add_divider", "add_quote_block", "add_callout",
    "add_severity_chip_run", "add_painpoint_card",
    "make_brand_reference_docx", "DESIGN_SYSTEM_PROMPT",
]
